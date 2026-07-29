# -*- coding: utf-8 -*-
import zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Emu
import copy
from .config import input_dir, output_dir, template_dir, weekly_report_path
from .parser import load_kegiatan_tambahan, parse_detail_github, parse_weekly_report

FONT = "Arial"
SIZE_BODY    = None   # inherit from Normal style
SIZE_HEADING = Pt(14)
SIZE_LAMPIRAN     = Pt(58)
SIZE_TABLE_CELL   = Pt(10)
LINE_SPACING      = 1.5   # 1.5 lines

def _set_line_spacing(pf, spacing=LINE_SPACING):
    pf.line_spacing = spacing

def _run(p, text, bold=False, size=None, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = size
    return r

def enable_update_fields(doc):
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings.element.append(update)


def add_page_numbers(doc):
    """Insert a centered PAGE field in the footer of every section.
    Cover page (first page of first section) is excluded via different_first_page.
    """
    for i, section in enumerate(doc.sections):
        if i == 0:
            # Cover: first page gets no page number
            section.different_first_page_header_footer = True
            # first_page_footer stays empty — no action needed

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_spacing(fp.paragraph_format)
        r = fp.add_run()
        r.font.name = FONT
        # PAGE field: begin
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r._r.append(fld_begin)
        # instruction
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        r._r.append(instr)
        # PAGE field: end
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r._r.append(fld_end)


def add_toc(doc):
    """Inject a Word TOC field that auto-populates when opened in Word."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    _set_line_spacing(p.paragraph_format)
    r = p.add_run()
    # fldChar begin
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fld_begin)
    # instrText
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r._r.append(instr)
    # fldChar separate
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r._r.append(fld_sep)
    # fldChar end
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_end)

def add_body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = align
    _set_line_spacing(p.paragraph_format)
    _run(p, text)
    return p

def add_module_heading(doc, text):
    """Module/BAB heading — uses Heading 2 style so TOC picks it up."""
    p = doc.add_paragraph(style="Heading 2")
    _set_line_spacing(p.paragraph_format)
    # Clear existing runs from style and add our own
    for run in p.runs:
        run.text = ""
    r = p.add_run(text)
    r.font.name = FONT
    r.bold = True
    r.font.size = SIZE_HEADING
    return p

def add_activity_name(doc, text):
    """Sub-activity name — uses Heading 3 style so TOC picks it up."""
    p = doc.add_paragraph(style="Heading 3")
    _set_line_spacing(p.paragraph_format)
    p.paragraph_format.left_indent = Pt(14)
    for run in p.runs:
        run.text = ""
    r = p.add_run(text)
    r.font.name = FONT
    r.bold = False
    r.font.size = SIZE_HEADING
    return p

def add_lampiran_heading(doc):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p.paragraph_format)
    _run(p, "LAMPIRAN", bold=True, size=SIZE_LAMPIRAN)
    return p

# -------------------------------------------------------------------
# Activity table (2-col: label | content)
# -------------------------------------------------------------------

def _cell_run(cell, text, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.clear()
    p = cell.paragraphs[0]
    _set_line_spacing(p.paragraph_format)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SIZE_TABLE_CELL
    r.bold = bold
    return r

def _add_cell_para(cell, text, bold=False):
    p = cell.add_paragraph()
    _set_line_spacing(p.paragraph_format)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SIZE_TABLE_CELL
    r.bold = bold
    return p

def add_activity_table(doc, activity):
    COL0_W = Emu(1028700)   # ~0.81 in
    COL1_W = Emu(5382260)   # ~4.24 in

    deskripsi_text = activity.get("deskripsi", "")
    files = activity.get("files", [])
    utama = activity.get("utama", [])
    manfaat = activity.get("manfaat", [])

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    for row in table.rows:
        row.cells[0].width = COL0_W
        row.cells[1].width = COL1_W

    _cell_run(table.rows[0].cells[0], "Deskripsi Pekerjaan", bold=True)
    c1 = table.rows[0].cells[1]
    _cell_run(c1, deskripsi_text, bold=False)

    if files or utama or manfaat:
        _add_cell_para(c1, "", bold=False)
        _add_cell_para(c1, "Detail Perubahan", bold=True)
        if files:
            _add_cell_para(c1, "File yang Diubah:", bold=True)
            for f in files:
                _add_cell_para(c1, f)
        if utama:
            _add_cell_para(c1, "Perubahan Utama:", bold=True)
            for u in utama:
                _add_cell_para(c1, u)
        if manfaat:
            _add_cell_para(c1, "Manfaat:", bold=True)
            for m in manfaat:
                _add_cell_para(c1, m)

    _cell_run(table.rows[1].cells[0], "Dokumentasi", bold=True)
    docs = activity.get("docs", [])
    c1_dok = table.rows[1].cells[1]
    if docs:
        _cell_run(c1_dok, docs[0], bold=False)
        for d in docs[1:]:
            _add_cell_para(c1_dok, d)
    else:
        _cell_run(c1_dok, "-", bold=False)

    doc.add_paragraph()  # spacer after table

# -------------------------------------------------------------------
# Builders for BAB II and Lampiran
# -------------------------------------------------------------------
def build_bab2(doc, modules):
    for mod in modules:
        add_module_heading(doc, mod["title"])
        if mod["intro"]:
            add_body(doc, mod["intro"])
        for act in mod["activities"]:
            add_activity_name(doc, act["subtitle"])
            if act["prolog"]:
                add_body(doc, act["prolog"])
            add_activity_table(doc, act)

def build_kegiatan_tambahan(doc, kegiatan):
    if not kegiatan:
        return
    add_module_heading(doc, "Kegiatan Tambahan")
    for act in kegiatan:
        add_activity_name(doc, act["judul"])
        add_activity_table(doc, act)

def ensure_media(media_dir, template_path):
    if media_dir.exists() and any(media_dir.iterdir()):
        return
    media_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(template_path) as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                (media_dir / name.split("/")[-1]).write_bytes(z.read(name))

def add_image_if_exists(doc, media_dir, name, width=Inches(6.0)):
    path = media_dir / name
    if path.exists():
        doc.add_picture(str(path), width=width)
        return True
    return False

def build_manajemen_kode(doc, config, prs, media_dir):
    add_module_heading(doc, "Manajemen Kode Sumber")
    bulan = config.get("bulan", "")
    tahun = config.get("tahun", "")
    repo = config["github"]["repo"]
    author = config["github"]["author"]
    
    add_activity_name(doc, f"Sinkronisasi Kode Sumber pada Repositori Github {bulan} {tahun}")
    add_body(doc, f"Pada bulan {bulan} {tahun} dilakukan {len(prs)} pull request (author: {author}) di https://github.com/{repo}")
    
    if not add_image_if_exists(doc, media_dir, "image3.jpeg", width=Inches(6.0)):
        add_body(doc, "[Screenshot GitHub terlampir]")
    
    add_activity_name(doc, f"Sinkronisasi Kode Sumber pada Server Development {bulan} {tahun}")
    add_body(doc, "Sinkronisasi kode sumber API ke server development setelah perubahan di-merge.")
    if not add_image_if_exists(doc, media_dir, "image4.jpeg", width=Inches(6.0)):
        add_body(doc, "[Screenshot server terlampir]")

def build_lampiran(doc, config, prs, weekly_rows, media_dir):
    bulan = config.get("bulan", "")
    tahun = config.get("tahun", "")
    # Page break: LAMPIRAN cover is already in the template before the marker,
    # so push actual content to the next page.
    pb = doc.add_paragraph()
    pb.style = doc.styles["Normal"]
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    pb._p.append(br)

    # Weekly Report
    add_module_heading(doc, f"Lampiran 1 – Weekly Report {bulan} {tahun}")
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Weekly Cat", "Date Range", "Date", "Activity", "Output", "User", "Related Doc"]):
        _cell_run(tbl.rows[0].cells[i], h, bold=True)
    for row in weekly_rows:
        cells = tbl.add_row().cells
        _cell_run(cells[0], row.get("minggu", ""))
        _cell_run(cells[1], row.get("date_range", ""))
        _cell_run(cells[2], row.get("tanggal", ""))
        _cell_run(cells[3], row.get("aktivitas", ""))
        _cell_run(cells[4], row.get("output", ""))
        _cell_run(cells[5], row.get("user", ""))
        _cell_run(cells[6], row.get("dokumentasi", ""))
    
    # Code source & PR
    doc.add_page_break()
    add_module_heading(doc, "Lampiran 2 – Kode Sumber & PR List")
    add_body(doc, f"Bahasa Pemrograman\t: {config['tech']['language']}")
    add_body(doc, f"Framework\t\t\t: {config['tech']['framework']}")
    add_body(doc, f"Repositori\t\t\t: {config['tech']['repo_url']}")
    
    add_body(doc, f"Daftar Pull Request Bulan {bulan} {tahun} (Author: {config['github']['author']}) — Total: {len(prs)} PR")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["No", "PR #", "Judul", "Status"]):
        _cell_run(table.rows[0].cells[i], h, bold=True)
    for idx, pr in enumerate(prs, 1):
        row = table.add_row().cells
        _cell_run(row[0], str(idx))
        _cell_run(row[1], str(pr["number"]))
        _cell_run(row[2], pr["title"])
        _cell_run(row[3], pr["state"].upper())

    if not add_image_if_exists(doc, media_dir, "image12.jpeg"):
        doc.add_paragraph("[Screenshot GitHub PR terlampir]")

# -------------------------------------------------------------------
# The Core Splicing engine
# -------------------------------------------------------------------
def splice_elements(marker_p, builder_funcs, args_list):
    temp_doc = Document()
    for func, args in zip(builder_funcs, args_list):
        func(temp_doc, *args)
    
    curr = marker_p._element
    for el in temp_doc.element.body:
        if el.tag.endswith('sectPr'): continue
        new_el = copy.deepcopy(el)
        curr.addnext(new_el)
        curr = new_el
    marker_p._element.getparent().remove(marker_p._element)

def build_report(config, prs):
    period   = config["period"]
    inp, tpl = input_dir(period), template_dir()
    media_dir    = tpl / "media"
    template_path = tpl / "laporan_template.docx"
    output_path  = output_dir(period) / config["output_filename"]

    ensure_media(media_dir, template_path)
    modules      = parse_detail_github(inp / "detail_github.md") if (inp / "detail_github.md").exists() else []
    weekly_rows  = parse_weekly_report(
        weekly_report_path(period), 
        config.get("weekly_user_filter", "Lutfi"),
        repo_filter=config.get("github", {}).get("repo")
    )
    kegiatan     = load_kegiatan_tambahan(inp / "kegiatan_tambahan.json")

    doc = Document(str(template_path))
    enable_update_fields(doc)

    # Simple placeholder replacement on the original doc template
    bulan_up = config.get('bulan_up', config['bulan'].upper())
    replacements = {
        'Mei 2026': f"{config['bulan']} {config['tahun']}",
        'Mei2026': f"{config['bulan']} {config['tahun']}",
        'MEI 2026': f"{bulan_up} {config['tahun']}",
        '{{ BAB1_PENCAPAIAN }}': config.get('bab1', {}).get('pencapaian', ''),
        '{{ BAB1_TANTANGAN }}': config.get('bab1', {}).get('tantangan', ''),
    }
    # BAB 3 Kesimpulan is a list, join them with newlines
    kesimpulan = "\n".join(config.get('bab3', {}).get('kesimpulan_items', []))
    replacements['{{ BAB3_KESIMPULAN }}'] = kesimpulan
    for p in doc.paragraphs:
        for old_t, new_t in replacements.items():
            if old_t in p.text:
                for r in p.runs:
                    if old_t in r.text:
                        r.text = r.text.replace(old_t, new_t)
    
    toc_p = None
    bab2_p = None
    lampiran_p = None
    kata_pengantar_p = None
    for p in doc.paragraphs:
        if '{{ TOC }}' in p.text: toc_p = p
        if '{{ BAB2 }}' in p.text: bab2_p = p
        if '{{ LAMPIRAN }}' in p.text: lampiran_p = p
        if 'Kata Pengantar' == p.text.strip() and kata_pengantar_p is None:
            kata_pengantar_p = p

    # Inject TOC: if explicit marker found use it, else inject before Kata Pengantar
    if toc_p:
        splice_elements(toc_p, [add_toc], [()])
    elif kata_pengantar_p is not None:
        curr = kata_pengantar_p._element
        temp_doc = Document()
        # DAFTAR ISI heading + TOC field — NO leading page break,
        # template already provides natural page boundary before Kata Pengantar
        # Reuse the default empty paragraph Document() always creates,
        # instead of add_paragraph() which would create a second blank para
        # that becomes a blank page when injected.
        h_p = temp_doc.paragraphs[0] if temp_doc.paragraphs else temp_doc.add_paragraph()
        h_p.style = temp_doc.styles["Normal"]
        h_r = h_p.add_run("DAFTAR ISI")
        h_r.font.name = FONT
        h_r.bold = True
        h_r.font.size = SIZE_HEADING
        _set_line_spacing(h_p.paragraph_format)
        # TOC field
        add_toc(temp_doc)
        # Page break AFTER TOC so Kata Pengantar starts on new page
        pb2_p = temp_doc.add_paragraph()
        pb2_p.style = temp_doc.styles["Normal"]
        br2 = OxmlElement("w:br")
        br2.set(qn("w:type"), "page")
        pb2_p._p.append(br2)
        # Insert before kata_pengantar_p
        for el in temp_doc.element.body:
            if el.tag.endswith('sectPr'): continue
            new_el = copy.deepcopy(el)
            curr.addprevious(new_el)


    if bab2_p:
        splice_elements(bab2_p, [build_bab2, build_kegiatan_tambahan, build_manajemen_kode], [(modules,), (kegiatan,), (config, prs, media_dir)])

    if lampiran_p:
        splice_elements(lampiran_p, [build_lampiran], [(config, prs, weekly_rows, media_dir)])

    add_page_numbers(doc)
    doc.save(str(output_path))
    return output_path

