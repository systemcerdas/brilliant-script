# -*- coding: utf-8 -*-
import zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Emu
from .config import input_dir, output_dir, template_dir, weekly_report_path
from .parser import load_kegiatan_tambahan, parse_detail_github, parse_weekly_report

FONT = "Arial"
SIZE_BODY    = None   # inherit from Normal style (Times New Roman 12, overridden to Arial via run)
SIZE_HEADING = Pt(14)
SIZE_COVER_TITLE  = Pt(16)
SIZE_COVER_MONTH  = Pt(20)
SIZE_COVER_NAME   = Pt(14)
SIZE_LAMPIRAN     = Pt(58)
SIZE_TABLE_CELL   = Pt(10)
LINE_SPACING      = 1.5   # 1.5 lines

# -------------------------------------------------------------------
# Low-level helpers
# -------------------------------------------------------------------

def _set_line_spacing(pf, spacing=LINE_SPACING):
    pf.line_spacing = spacing

def _run(p, text, bold=False, size=None, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = size
    return r

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag.split("}")[-1] != "sectPr":
            body.remove(child)

def enable_update_fields(doc):
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings.element.append(update)

# -------------------------------------------------------------------
# Paragraph factories
# -------------------------------------------------------------------

def _blank(doc):
    p = doc.add_paragraph()
    _set_line_spacing(p.paragraph_format)
    return p

def add_body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Normal body paragraph – Arial, inherit size, justified, 1.5 spacing."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = align
    _set_line_spacing(p.paragraph_format)
    _run(p, text)
    return p

def add_center(doc, text, bold=False, size=None):
    """Centred paragraph for cover page."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p.paragraph_format)
    _run(p, text, bold=bold, size=size)
    return p

def add_bab_heading(doc, text):
    """BAB I / BAB II / BAB III / Kata Pengantar  — Normal, 14pt Bold, no indent."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    _set_line_spacing(p.paragraph_format)
    _run(p, text, bold=True, size=SIZE_HEADING)
    return p

def add_section_heading(doc, text):
    """Sub-section headings: Latar Belakang, Kesimpulan, Saran, etc.
    Uses List Paragraph style, Arial 14pt Bold."""
    p = doc.add_paragraph(style="List Paragraph")
    _set_line_spacing(p.paragraph_format)
    p.paragraph_format.left_indent = Pt(0)
    _run(p, text, bold=True, size=SIZE_HEADING)
    return p

def add_module_heading(doc, text):
    """Module title in BAB II — same as BAB heading, 14pt Bold."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    _set_line_spacing(p.paragraph_format)
    _run(p, text, bold=True, size=SIZE_HEADING)
    return p

def add_activity_name(doc, text):
    """Sub-activity name — List Paragraph, Arial inherit, 14pt indented."""
    p = doc.add_paragraph(style="List Paragraph")
    _set_line_spacing(p.paragraph_format)
    p.paragraph_format.left_indent = Pt(14)
    _run(p, text, bold=False, size=SIZE_HEADING)
    return p

def add_lampiran_heading(doc):
    """LAMPIRAN — centred, Arial 58pt Bold."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p.paragraph_format)
    _run(p, "LAMPIRAN", bold=True, size=SIZE_LAMPIRAN)
    return p

def add_toc(doc):
    add_bab_heading(doc, "DAFTAR ISI")
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    _set_line_spacing(p.paragraph_format)
    r = p.add_run()
    for el_type, text in [("begin", None), ("instr", ' TOC \\o "1-3" \\h \\z \\u '), ("separate", None), ("end", None)]:
        if el_type == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            r._r.append(instr)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), el_type)
            r._r.append(fld)
    add_body(doc, "(Klik kanan Daftar Isi -> Update Field jika nomor halaman belum muncul)")

# -------------------------------------------------------------------
# Activity table (2-col: label | content)
# -------------------------------------------------------------------

def _cell_run(cell, text, bold=False):
    """Set cell text with Arial 10pt."""
    # Clear existing
    for p in cell.paragraphs:
        for r in p.runs:
            r.clear()
    p = cell.paragraphs[0]
    _set_line_spacing(p.paragraph_format)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SIZE_TABLE_CELL
    r.bold = bold
    return r

def _add_cell_para(cell, text, bold=False):
    """Add an additional paragraph inside a cell."""
    p = cell.add_paragraph()
    _set_line_spacing(p.paragraph_format)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SIZE_TABLE_CELL
    r.bold = bold
    return p

def add_activity_table(doc, activity):
    """Render one activity as a 2-col table matching the original format.

    Col 0 (narrow ~1.1in): row labels (Deskripsi Pekerjaan, Dokumentasi)
    Col 1 (wide ~4.3in):   content (deskripsi text, files, docs, etc.)
    """
    COL0_W = Emu(1028700)   # ~0.81 in
    COL1_W = Emu(5382260)   # ~4.24 in

    # Build content for col1 row 0 (Deskripsi Pekerjaan)
    deskripsi_text = activity.get("deskripsi", "")
    files = activity.get("files", [])
    utama = activity.get("utama", [])
    manfaat = activity.get("manfaat", [])

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    # Set column widths
    for row in table.rows:
        row.cells[0].width = COL0_W
        row.cells[1].width = COL1_W

    # ROW 0 — Deskripsi Pekerjaan
    _cell_run(table.rows[0].cells[0], "Deskripsi Pekerjaan", bold=True)
    c1 = table.rows[0].cells[1]
    _cell_run(c1, deskripsi_text, bold=False)

    if files or utama or manfaat:
        _add_cell_para(c1, "", bold=False)  # spacer
        _add_cell_para(c1, "Detail Perubahan", bold=True)
        if files:
            _add_cell_para(c1, "File yang Diubah:", bold=True)
            for f in files:
                _add_cell_para(c1, f)
        if utama:
            _add_cell_para(c1, "Perubahan Utama:", bold=True)
            for u in utama:
                _add_cell_para(c1, u)
        if manfaat:
            _add_cell_para(c1, "Manfaat:", bold=True)
            for m in manfaat:
                _add_cell_para(c1, m)

    # ROW 1 — Dokumentasi
    _cell_run(table.rows[1].cells[0], "Dokumentasi", bold=True)
    docs = activity.get("docs", [])
    c1_dok = table.rows[1].cells[1]
    if docs:
        _cell_run(c1_dok, docs[0], bold=False)
        for d in docs[1:]:
            _add_cell_para(c1_dok, d)
    else:
        _cell_run(c1_dok, "-", bold=False)

    doc.add_paragraph()  # spacer after table

# -------------------------------------------------------------------
# Images
# -------------------------------------------------------------------

def ensure_media(media_dir, template_path):
    if media_dir.exists() and any(media_dir.iterdir()):
        return
    media_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(template_path) as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                (media_dir / name.split("/")[-1]).write_bytes(z.read(name))

def add_image_if_exists(doc, media_dir, name, width=Inches(6.0)):
    path = media_dir / name
    if path.exists():
        doc.add_picture(str(path), width=width)
        return True
    return False

# -------------------------------------------------------------------
# PR Table (Lampiran)
# -------------------------------------------------------------------

def add_pr_table(doc, prs):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["No", "PR #", "Judul", "Status"]
    for i, h in enumerate(headers):
        _cell_run(table.rows[0].cells[i], h, bold=True)
    for idx, pr in enumerate(prs, 1):
        row = table.add_row().cells
        _cell_run(row[0], str(idx))
        _cell_run(row[1], str(pr["number"]))
        _cell_run(row[2], pr["title"])
        _cell_run(row[3], pr["state"].upper())

# -------------------------------------------------------------------
# Weekly table (Lampiran 1)
# -------------------------------------------------------------------

def add_weekly_table(doc, weekly_rows):
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    headers = ["Weekly Cat", "Date Range", "Date", "Activity", "Output", "User", "Related Doc"]
    for i, h in enumerate(headers):
        _cell_run(tbl.rows[0].cells[i], h, bold=True)
    for row in weekly_rows:
        cells = tbl.add_row().cells
        _cell_run(cells[0], row.get("minggu", ""))
        _cell_run(cells[1], row.get("date_range", ""))
        _cell_run(cells[2], row.get("tanggal", ""))
        _cell_run(cells[3], row.get("aktivitas", ""))
        _cell_run(cells[4], row.get("output", ""))
        _cell_run(cells[5], row.get("user", ""))
        _cell_run(cells[6], row.get("dokumentasi", ""))

# -------------------------------------------------------------------
# BAB I
# -------------------------------------------------------------------

def build_report(config, prs):
    period   = config["period"]
    inp, tpl = input_dir(period), template_dir()
    media_dir    = tpl / "media"
    template_path = tpl / "laporan_template.docx"
    output_path  = output_dir(period) / config["output_filename"]

    ensure_media(media_dir, template_path)
    modules      = parse_detail_github(inp / "detail_github.md") if (inp / "detail_github.md").exists() else []
    weekly_rows  = parse_weekly_report(weekly_report_path(period), config.get("weekly_user_filter", "Lutfi"))
    kegiatan     = load_kegiatan_tambahan(inp / "kegiatan_tambahan.json")

    doc = Document(str(template_path))
    clear_body(doc)
    enable_update_fields(doc)

    bulan_up = config.get("bulan_up", config["bulan"].upper())
    nama     = config.get("nama", "Lutfi Ihsan")
    jabatan  = config.get("jabatan", "Tenaga Teknis Implementasi Logika Sistem")
    tahun    = config["tahun"]

    # ---- COVER PAGE ----
    for _ in range(3):
        _blank(doc)
    add_center(doc, "LAPORAN KEMAJUAN", bold=True, size=SIZE_COVER_TITLE)
    add_center(doc, "TENAGA TEKNIS IMPLEMENTASI LOGIKA SISTEM", bold=True, size=SIZE_COVER_TITLE)
    add_center(doc, "DALAM RANGKA PENGELOLAAN DATA PENGAWASAN SDKP", bold=True, size=SIZE_COVER_TITLE)
    _blank(doc)
    _blank(doc)
    add_center(doc, bulan_up, bold=True, size=SIZE_COVER_MONTH)
    _blank(doc)
    _blank(doc)
    add_center(doc, nama, bold=True, size=SIZE_COVER_NAME)
    add_center(doc, jabatan, bold=True, size=SIZE_COVER_NAME)
    _blank(doc)
    _blank(doc)
    add_center(doc, "DIREKTORAT JENDERAL PENGAWASAN SUMBER DAYA", bold=True, size=SIZE_COVER_NAME)
    add_center(doc, "KELAUTAN DAN PERIKANAN", bold=True, size=SIZE_COVER_NAME)
    add_center(doc, "KEMENTERIAN KELAUTAN DAN PERIKANAN", bold=True, size=SIZE_COVER_TITLE)
    add_center(doc, str(tahun), bold=True, size=SIZE_COVER_TITLE)

    # ---- DAFTAR ISI ----
    doc.add_page_break()
    add_toc(doc)

    # ---- KATA PENGANTAR ----
    doc.add_page_break()
    add_bab_heading(doc, "Kata Pengantar")
    _blank(doc)
    add_body(doc, config.get("kata_pengantar_1",
        "Kami ingin memulai laporan ini dengan ucapan terima kasih dan puji syukur kepada Tuhan Yang Maha Esa atas segala berkat dan kemudahan yang telah diberikan kepada kami dalam menjalankan tugas sebagai Tenaga Teknis Implementasi Logika Sistem."))
    add_body(doc, config.get("kata_pengantar_2",
        f"Pada bulan ini, kami sebagai Tenaga Teknis Implementasi Logika Sistem telah berupaya keras untuk memantau, menyusun, mengimplementasikan, dan memelihara sistem pengawasan SDKP."))
    _blank(doc)
    add_body(doc, config.get("kata_pengantar_3",
        "Kami berharap bahwa laporan ini dapat memberikan manfaat bagi pembaca dan pihak yang terkait dalam pengawasan sumber daya kelautan dan perikanan."))
    _blank(doc)
    add_body(doc, config.get("kata_pengantar_4",
        "Kami mengucapkan terima kasih atas kesempatan ini dan menerima dengan terbuka segala masukan dan saran yang membangun guna meningkatkan kualitas kerja kami ke depannya."))
    _blank(doc)

    # Tanda tangan kanan
    p_ttd = doc.add_paragraph()
    p_ttd.style = doc.styles["Normal"]
    _set_line_spacing(p_ttd.paragraph_format)
    p_ttd.paragraph_format.left_indent = Pt(354)
    _run(p_ttd, f"Jakarta,    {config['bulan']} {tahun} ")
    p_kami = doc.add_paragraph()
    p_kami.style = doc.styles["Normal"]
    _set_line_spacing(p_kami.paragraph_format)
    p_kami.paragraph_format.left_indent = Pt(354)
    _run(p_kami, "Hormat Kami,")
    for _ in range(3):
        p_sp = doc.add_paragraph()
        p_sp.style = doc.styles["Normal"]
        _set_line_spacing(p_sp.paragraph_format)
        p_sp.paragraph_format.left_indent = Pt(396)
    p_nama = doc.add_paragraph()
    p_nama.style = doc.styles["Normal"]
    _set_line_spacing(p_nama.paragraph_format)
    p_nama.paragraph_format.left_indent = Pt(354)
    _run(p_nama, nama)

    # ---- BAB I ----
    doc.add_page_break()
    build_bab1(doc, config)

    # ---- BAB II ----
    doc.add_page_break()
    build_bab2(doc, modules)
    build_kegiatan_tambahan(doc, kegiatan)
    build_manajemen_kode(doc, config, prs, media_dir)

    # ---- BAB III ----
    doc.add_page_break()
    build_bab3(doc, config, prs)

    # ---- LAMPIRAN ----
    build_lampiran(doc, config, prs, weekly_rows, media_dir)

    doc.save(str(output_path))
    return output_path
