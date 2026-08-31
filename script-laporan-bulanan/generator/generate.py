# -*- coding: utf-8 -*-
import argparse
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from lib.config import input_dir, load_config
from lib.docx_builder import build_report, build_diff_document
from lib.github import load_or_fetch_prs, load_or_fetch_diffs
from docx import Document
from docxcompose.composer import Composer
import time


def main():
    parser = argparse.ArgumentParser(description="Generate Laporan Bulanan")
    parser.add_argument("period", help="Format YYYYMM (contoh: 202607)")
    parser.add_argument("--fetch-prs", action="store_true", help="Paksa ambil ulang PR dari GitHub")
    args = parser.parse_args()

    config = load_config(args.period)
    if args.fetch_prs:
        config.setdefault("github", {})["fetch_always"] = True

    prs = load_or_fetch_prs(config, input_dir(args.period) / "prs.json")
    diffs = load_or_fetch_diffs(config, prs, input_dir(args.period) / "prs_diff.json")
    
    # 1. Build Main Document (BAB 1 - 3)
    try:
        output_main = build_report(config, prs, diffs)
    except PermissionError:
        print("File is locked, saving as _v2")
        config["output_filename"] = config["output_filename"].replace(".docx", " _v2.docx")
        output_main = build_report(config, prs, diffs)
        
    print("Membangun dokumen Lampiran 2 (Diff)...")
    try:
        output_diff = build_diff_document(config, prs, diffs)
    except PermissionError:
        print("File Lampiran Diff is locked, appending _v2")
        if "_v2" not in config["output_filename"]:
            config["output_filename"] = config["output_filename"].replace(".docx", " _v2.docx")
            output_main = build_report(config, prs, diffs) # Re-build main to match filename
        output_diff = build_diff_document(config, prs, diffs)
        
    print("Menggabungkan Lampiran 1 (Manual) dan Lampiran 2...")
    master = Document(output_main)
    
    # 2. Append Lampiran 1 (Manual Weekly Report Landscape)
    manual_weekly_path = input_dir(args.period) / f"{args.period}_Program dan Data Weekly Report.docx"
    if manual_weekly_path.exists():
        doc_weekly = Document(str(manual_weekly_path))
        
        # Force a section break and copy landscape dimensions to prevent docxcompose from stripping it
        from docx.enum.section import WD_SECTION
        new_sect = master.add_section(WD_SECTION.NEW_PAGE)
        if doc_weekly.sections:
            src_sect = doc_weekly.sections[0]
            new_sect.orientation = src_sect.orientation
            new_sect.page_width = src_sect.page_width
            new_sect.page_height = src_sect.page_height
            
        # Add heading for Lampiran 1
        from lib.docx_builder import add_module_heading
        add_module_heading(master, f"Lampiran 1 – Weekly Report {config.get('bulan', '')} {config.get('tahun', '')}")
            
        composer = Composer(master)
        composer.append(doc_weekly)
    else:
        print(f"WARNING: File manual weekly report tidak ditemukan di {manual_weekly_path}")
        composer = Composer(master)
        
    # 3. Append Lampiran 2 (Code Diff)
    doc_diff = Document(str(output_diff))
    # We want Lampiran 2 to be Portrait again.
    # Add a section break to return to portrait.
    from docx.enum.section import WD_SECTION
    new_sect_diff = master.add_section(WD_SECTION.NEW_PAGE)
    if doc_diff.sections:
        src_sect_diff = doc_diff.sections[0]
        new_sect_diff.orientation = src_sect_diff.orientation
        new_sect_diff.page_width = src_sect_diff.page_width
        new_sect_diff.page_height = src_sect_diff.page_height
        
    composer.append(doc_diff)
    
    # 4. Save Final
    composer.save(str(output_main))
    
    # Clean up temp diff file
    if Path(output_diff).exists():
        Path(output_diff).unlink()
    
    print("Selesai:", output_main)
    print("Total PR:", len(prs))
    print("Buka di Word -> klik kanan Daftar Isi -> Update Field untuk nomor halaman.")


if __name__ == "__main__":
    main()
