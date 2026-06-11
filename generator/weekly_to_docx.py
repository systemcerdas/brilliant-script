# -*- coding: utf-8 -*-
"""Generate weekly_report.docx dari weekly_report.md"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from lib.config import input_dir, project_root


def clean_md_links(text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text.replace("\\_", "_").strip()


def parse_md_rows(md_path):
    rows = []
    for line in md_path.read_text(encoding="utf-8").splitlines():
        if not line.startswith("|") or "Weekly" in line or "-----" in line:
            continue
        cols = [clean_md_links(c.strip()) for c in line.split("|")[1:-1]]
        if len(cols) < 7:
            cols += [""] * (7 - len(cols))
        rows.append(cols[:7])
    return rows


def set_cell_text(cell, text, font_size=9):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(font_size)
            run.font.name = "Calibri"


def build_weekly_docx(period):
    root = project_root()
    inp = input_dir(period)
    md_path = inp / "weekly_report.md"
    template = root / "template" / "weekly_report_template.docx"
    output = inp / "weekly_report.docx"

    if not md_path.exists():
        raise FileNotFoundError(md_path)
    if not template.exists():
        raise FileNotFoundError(template)

    rows = parse_md_rows(md_path)
    doc = Document(str(template))
    table = doc.tables[0]

    while len(table.rows) - 1 < len(rows):
        table.add_row()
    while len(table.rows) - 1 > len(rows) and len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for i, cols in enumerate(rows, start=1):
        for j, val in enumerate(cols):
            set_cell_text(table.rows[i].cells[j], val)

    doc.save(str(output))
    return output, len(rows)


def main():
    parser = argparse.ArgumentParser(description="Generate weekly_report.docx dari markdown")
    parser.add_argument("period", help="YYYYMM contoh 202605")
    args = parser.parse_args()

    output, count = build_weekly_docx(args.period)
    print("Selesai:", output)
    print("Baris:", count)


if __name__ == "__main__":
    main()
